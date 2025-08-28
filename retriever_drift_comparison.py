import json
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL

def load_analysis_data(filename):
    """Load retriever drift analysis data from JSON file"""
    with open(filename, 'r') as f:
        return json.load(f)

def create_comparison_document():
    """Create Word document with comparison tables"""
    # Load data from both files
    data_11_files = load_analysis_data('retriever_drift_analysis.json')
    data_2_files = load_analysis_data('retriever_drift_analysis(2).json')
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Retriever Drift Analysis Comparison', 0)
    title.alignment = 1  # Center alignment
    
    # Add description
    doc.add_paragraph('This document compares retriever drift analysis results between two datasets: one with 11 files and another with 2 files.')
    
    # Table 1: 11 Files Analysis
    doc.add_heading('Analysis Results - 11 Files Dataset', level=1)
    doc.add_paragraph(f'Total chunks: {data_11_files["nsf_data_summary"]["total_chunks"]}')
    doc.add_paragraph(f'Unique sources: {data_11_files["nsf_data_summary"]["unique_sources"]}')
    
    # Create table for 11 files
    table1 = doc.add_table(rows=1, cols=8)
    table1.style = 'Table Grid'
    
    # Add headers
    headers = ['Query', 'ID Stability (%)', 'ID Drift (%)', 'Source Stability (%)', 'Source Drift (%)', 'Content Stability (%)', 'Content Drift (%)', 'Successful Runs']
    for i, header in enumerate(headers):
        cell = table1.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add data rows
    for query, metrics in data_11_files['improved_drift_analysis'].items():
        row = table1.add_row()
        row.cells[0].text = query
        row.cells[1].text = str(metrics['id_stability'])
        row.cells[2].text = str(metrics['id_drift'])
        row.cells[3].text = str(metrics['source_stability'])
        row.cells[4].text = str(metrics['source_drift'])
        row.cells[5].text = str(metrics['content_stability'])
        row.cells[6].text = str(metrics['content_drift'])
        row.cells[7].text = str(metrics['successful_runs'])
        
        # Center align all cells
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Table 2: 2 Files Analysis
    doc.add_heading('Analysis Results - 2 Files Dataset', level=1)
    doc.add_paragraph(f'Total chunks: {data_2_files["nsf_data_summary"]["total_chunks"]}')
    doc.add_paragraph(f'Unique sources: {data_2_files["nsf_data_summary"]["unique_sources"]}')
    
    # Create table for 2 files
    table2 = doc.add_table(rows=1, cols=8)
    table2.style = 'Table Grid'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table2.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add data rows
    for query, metrics in data_2_files['improved_drift_analysis'].items():
        row = table2.add_row()
        row.cells[0].text = query
        row.cells[1].text = str(metrics['id_stability'])
        row.cells[2].text = str(metrics['id_drift'])
        row.cells[3].text = str(metrics['source_stability'])
        row.cells[4].text = str(metrics['source_drift'])
        row.cells[5].text = str(metrics['content_stability'])
        row.cells[6].text = str(metrics['content_drift'])
        row.cells[7].text = str(metrics['successful_runs'])
        
        # Center align all cells
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add summary section
    doc.add_heading('Key Insights', level=1)
    
    insights = [
        "• Both datasets show 100% ID and Source stability, indicating perfect consistency in chunk and source selection",
        "• Content stability varies between datasets, showing some semantic variation in retrieved content",
        "• The 11-files dataset has higher content stability for NYU queries (81.37% vs 71.9%)",
        "• COSMOS queries maintain consistent content stability across both datasets (63.86%)",
        "• All runs were successful (5/5) in both analyses",
        "• The smaller dataset (2 files) shows slightly higher content drift, suggesting more semantic variation"
    ]
    
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    
    # Save document
    doc.save('retriever_drift_comparison.docx')
    print("Word document 'retriever_drift_comparison.docx' has been created successfully!")

if __name__ == "__main__":
    create_comparison_document()

