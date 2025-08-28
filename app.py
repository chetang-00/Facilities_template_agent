import streamlit as st
from src.pdf_ingest import ingest_pdfs
from src.generate import generate_enriched_response, get_section_labels_for_agency
from src.utils import get_llm
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from io import BytesIO
from docx import Document

# Page Config ────────────────────────────────
st.set_page_config(page_title="Grant Facilities Draft Generator", layout="wide")

# Utility Function ───────────────────────────
def build_full_draft(sections_dict, section_labels):
    final = ""
    for section_label in section_labels:
        section_text = sections_dict.get(section_label, "").strip()
        if section_text:
            final += f"\n\n## {section_label}\n\n{section_text}\n"
    return final.strip()

def generate_pdf(sections_dict, section_labels, sources):
    """Generate a PDF from the sections and sources"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1 
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20
    )
    normal_style = styles['Normal']
    
    # Add title
    story.append(Paragraph("Facilities Template", title_style))
    story.append(Spacer(1, 20))
    
    # Add sections
    for section_label in section_labels:
        section_text = sections_dict.get(section_label, "").strip()
        if section_text:
            story.append(Paragraph(section_label, heading_style))
            story.append(Paragraph(section_text, normal_style))
            story.append(Spacer(1, 12))
    

    if sources:
        story.append(Paragraph("Sources Used", heading_style))
        for source in sources:
            story.append(Paragraph(f"• {source}", normal_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_word_doc(sections_dict, section_labels, sources):
    """Generate a Word document from the sections and sources"""
    doc = Document()
    
    # Add sections (no title)
    for section_label in section_labels:
        section_text = sections_dict.get(section_label, "").strip()
        if section_text:
            # Add section heading
            doc.add_heading(section_label, level=1)
            
            # Add section content
            doc.add_paragraph(section_text)
            
            # Add some spacing
            doc.add_paragraph()
    
    # Add sources if available
    if sources:
        doc.add_heading("Sources Used", level=1)
        for source in sources:
            doc.add_paragraph(f"• {source}", style='List Bullet')
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Session State Init ─────────────────────────
if "conversation_chain" not in st.session_state:
    memory = ConversationBufferMemory()
    llm = get_llm()
    st.session_state.conversation_chain = ConversationChain(llm=llm, memory=memory)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "draft_generated" not in st.session_state:
    st.session_state.draft_generated = False

if "enriched_sections" not in st.session_state:
    st.session_state.enriched_sections = {}

if "final_draft" not in st.session_state:
    st.session_state.final_draft = ""

if "sources" not in st.session_state:
    st.session_state.sources = []

if "section_labels" not in st.session_state:
    st.session_state.section_labels = []

# undo functionality session state
if "section_edit_history" not in st.session_state:
    st.session_state.section_edit_history = {}

if "section_original_content" not in st.session_state:
    st.session_state.section_original_content = {}

if "section_edit_messages" not in st.session_state:
    st.session_state.section_edit_messages = {}

# Layout ─────────────────────────────────────
left, right = st.columns([1, 2])

# Form ─────────────────────────────────
with left:
    st.title("Grant Facilities Section Form")

    if st.button("Reindex PDFs in `/data` folder"):
        ingest_pdfs()
        st.success("Reindex complete!")
    
    st.markdown("### Select Files to Search:")
    selected_agency = st.radio(
        "Choose one file type by Sponsor:",
        options=["NSF", "NIH"],
        index=None,
        help="Select the file type to determine which sections to include"
    )

    # Get section labels based on selected agency
    selected_types = []
    if selected_agency == "NIH":
        selected_types.append("NIH")
    elif selected_agency == "NSF":
        selected_types.append("NSF")
    
    section_labels = get_section_labels_for_agency(selected_types)

        # Only show form if a file type is selected
    if selected_agency:
        st.markdown("### Fill in the details below:")
        
        user_inputs = {}
        with st.form("input_form"):
            for label in section_labels:
                # Add some CSS to ensure labels are visible
                st.markdown(f"<div style=' color: black;'>{label}</div>", unsafe_allow_html=True)
                user_inputs[label] = st.text_area(label, height=100, key=f"input_{label}", label_visibility="hidden")
            submitted = st.form_submit_button("Generate Section")
    else:
        
        user_inputs = {}
        submitted = False

    if submitted:
        if not selected_agency:
            st.error("Please select files to search before generating.")
        else:
            with st.spinner("Generating and validating..."):
                
                section_outputs, sources = generate_enriched_response(user_inputs, selected_types=selected_types)

            if not section_outputs:
                st.warning("Please fill out the form to generate your Facilities Template.")
            else:
                st.session_state.enriched_sections = section_outputs
                st.session_state.final_draft = build_full_draft(section_outputs, section_labels)
                st.session_state.sources = sources
                st.session_state.section_labels = section_labels  # Store section labels
                st.session_state.draft_generated = True
                
                for section, content in section_outputs.items():
                    st.session_state.section_original_content[section] = content
                    st.session_state.section_edit_history[section] = []

                st.session_state.conversation_chain.memory.chat_memory.clear()
                st.session_state.conversation_chain.memory.chat_memory.add_user_message(
                    "This is the final enriched document:\n" + st.session_state.final_draft
                )
                st.session_state.chat_history.clear()
                st.success("Draft complete!")

with right:
    if st.session_state.draft_generated:
        #st.markdown("### Download PDF")
        
        if "show_filename_input" not in st.session_state:
            st.session_state.show_filename_input = False
        
        # Generate PDF and Word documents
        pdf_buffer = generate_pdf(st.session_state.enriched_sections, st.session_state.section_labels, st.session_state.sources)
        word_buffer = generate_word_doc(st.session_state.enriched_sections, st.session_state.section_labels, st.session_state.sources)
        
        # Create download buttons
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("Download as PDF", use_container_width=True, type="primary"):
                st.session_state.show_filename_input = "pdf"
        with col3:
            if st.button("Download as Word", use_container_width=True, type="primary"):
                st.session_state.show_filename_input = "word"
        
        if st.session_state.show_filename_input:
            st.markdown("---")
            default_filename = "facilities_template"
            
            if st.session_state.show_filename_input == "pdf":
                filename = st.text_input(
                    "Enter filename (without .pdf extension):",
                    value=default_filename,
                    help="The PDF will be saved with this name"
                )
                
                # download PDF button
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_buffer.getvalue(),
                        file_name=f"{filename}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="secondary"
                    )
            
            elif st.session_state.show_filename_input == "word":
                filename = st.text_input(
                    "Enter filename (without .docx extension):",
                    value=default_filename,
                    help="The Word document will be saved with this name"
                )
                
                # download Word button
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.download_button(
                        label="Download Word Document",
                        data=word_buffer.getvalue(),
                        file_name=f"{filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="secondary"
                    )
        
        st.markdown("### Final Draft (Sections)")

        for section in st.session_state.section_labels:
            text = st.session_state.enriched_sections.get(section, "")

            if not text:
                continue

            st.markdown(f"## {section}")

            if f"{section}_edit_mode" not in st.session_state:
                st.session_state[f"{section}_edit_mode"] = False

            if st.session_state[f"{section}_edit_mode"]:
                edited_text = st.text_area(
                    f"Edit {section}",
                    value=text,
                    height=200,
                    key=f"{section}_textarea"
                )

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Save Edits"):
                        # Store the previous version in edit history
                        if section in st.session_state.enriched_sections:
                            st.session_state.section_edit_history[section].append(
                                st.session_state.enriched_sections[section]
                            )
                        
                        st.session_state.enriched_sections[section] = edited_text
                        st.session_state[f"{section}_edit_mode"] = False
                        
                        st.session_state.section_edit_messages[section] = True

                        st.session_state.final_draft = build_full_draft(
                            st.session_state.enriched_sections, st.session_state.section_labels
                        )
                        st.session_state.conversation_chain.memory.chat_memory.clear()
                        st.session_state.conversation_chain.memory.chat_memory.add_user_message(
                            "This is the final enriched document:\n" + st.session_state.final_draft
                        )
                        st.success(f"Saved edits for {section}.")

                with col2:
                    if st.button("Cancel Edits"):
                        st.session_state[f"{section}_edit_mode"] = False

            else:
                st.markdown(text)
                st.button(
                    f" Edit {section}",
                    key=f"{section}_edit_button",
                    on_click=lambda s=section: st.session_state.update({f"{s}_edit_mode": True})
                )
                
                # edit message and revert button if section was edited
                if st.session_state.section_edit_messages.get(section, False):
                    st.info(f"{section} has been edited")
                    
                    if section in st.session_state.section_original_content:
                        if st.button("Revert to original", key=f"{section}_revert_button"):
                            # Restore original content
                            st.session_state.enriched_sections[section] = st.session_state.section_original_content[section]
                            
                            st.session_state.section_edit_history[section] = []
                            st.session_state.section_edit_messages[section] = False
                            
                            # Update final draft
                            st.session_state.final_draft = build_full_draft(
                                st.session_state.enriched_sections, st.session_state.section_labels
                            )
                            
                            # Update conversation memory
                            st.session_state.conversation_chain.memory.chat_memory.clear()
                            st.session_state.conversation_chain.memory.chat_memory.add_user_message(
                                "This is the final enriched document:\n" + st.session_state.final_draft
                            )
                            
                            st.success(f"Reverted {section} to original")
                            st.rerun()

        # Sources ──────────────────────────────
        st.markdown("### Sources Used")
        for s in st.session_state.sources:
            if s.endswith(".pdf"):
                st.markdown(f"- {s}")
            elif s.startswith("http"):
                st.markdown(f"- [{s}]({s})")
            else:
                st.markdown(f"- {s}")

    # Follow-up Chat ──────────────────────────
    st.markdown("### Follow-up Chat")

    if followup := st.chat_input("Ask a follow-up question..."):
        if not st.session_state.draft_generated:
            bot_reply = "Please fill out the form to begin generating your Facilities Template."
        else:
            st.session_state.conversation_chain.memory.chat_memory.clear()
            st.session_state.conversation_chain.memory.chat_memory.add_user_message(
                "This is the final enriched document:\n" + st.session_state.final_draft
            )

            bot_reply = st.session_state.conversation_chain.predict(input=followup)

            st.session_state.conversation_chain.memory.chat_memory.add_user_message(followup)
            st.session_state.conversation_chain.memory.chat_memory.add_ai_message(bot_reply)

        st.session_state.chat_history.append((followup, bot_reply))

    for user_q, bot_a in st.session_state.chat_history:
        st.markdown(f"**You:** {user_q}")
        st.markdown(f"**Assistant:** {bot_a}")
